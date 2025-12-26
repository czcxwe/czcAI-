# -*- coding: utf-8 -*-
"""
md2docx.py
改进点：
- 强制 Pandoc 将 Markdown 管道表格按原样识别为 Word 表格（使用 gfm 表格语法支持）。
- 保留中文、数学（$...$, $$...$$）与行内换行，不吞段落。
- 对“表格单元格内的 $...$ 或 =... ”：
  - 若以 "=" 开头，保留为纯文本（Word 表格不支持计算），便于后续复制到 Excel。
  - 若以 $...$ 包裹，去掉 $ 包裹外壳并保留表达式文本（不转义为 Unicode 数学），避免 docx 里变成乱码或被合并。
实现方式：
- 先用 Pandoc 转 docx（启用 gfm+tex_math_dollars）。
- 然后用 python-docx 逐个遍历表格单元，做“公式文本清洁”：
  - 单元格为 $...$ 或 $$...$$ 的，去壳；
  - 单元格为 =... 的，设置等宽字体以示为公式。
说明：
- Word 表格无法执行 Excel 公式，本脚本只在视觉上保留公式文本；需要可计算请改为嵌入式 Excel（Windows+Office + win32com）。
用法：
python3 md2docx.py 2.md 2.docx
"""

import argparse
import os
import shutil
import subprocess
import sys
import tempfile
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def check_pandoc():
    return shutil.which("pandoc") is not None

def preprocess_markdown(in_path):
    """
    轻量预处理：
    - 在块元素之间确保一个空行，帮助 Pandoc 区分段落。
    - 对包含中文标点/公式且下一行非空的软换行，追加两个空格，强制硬换行。
    - 不改变表格语法。
    """
    with open(in_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    out_lines = []
    prev_blank = True

    def is_block_start(line):
        s = line.strip()
        return (
            s.startswith("#") or
            s.startswith(">") or
            s.startswith("```") or s.startswith("~~~") or
            s.startswith("$$") or
            s == ""
        )

    for i, line in enumerate(lines):
        cur = line.rstrip("\n")
        # 保持表格原样
        is_table_row = cur.strip().startswith("|")
        if is_block_start(cur) and not is_table_row:
            if not prev_blank and cur != "":
                out_lines.append("")
            out_lines.append(cur)
            prev_blank = (cur == "")
            continue

        # 对密集中文/公式行追加硬换行
        next_line = lines[i+1].rstrip("\n") if i+1 < len(lines) else ""
        dense = ("$" in cur) or ("：" in cur and len(cur) > 40) or ("," in cur and len(cur) > 80)
        if cur != "" and next_line != "" and not next_line.strip().startswith(("#", ">", "```", "~~~", "$$")):
            if dense and not cur.endswith("  "):
                cur = cur + "  "

        out_lines.append(cur)
        prev_blank = (cur == "")

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".md", mode="w", encoding="utf-8")
    tmp.write("\n".join(out_lines))
    tmp.close()
    return tmp.name

def convert_md_to_docx(input_md, output_docx, reference_doc=None):
    """
    使用 Pandoc 将 Markdown 转 DOCX：
    - 启用 gfm（GitHub 风格）以稳定识别管道表格为 Word 表格；
    - tex_math_dollars 以保留 $...$ 数学；
    - wrap=none 减少软换行合并。
    """
    cmd = [
        "pandoc",
        "--from", "gfm+tex_math_dollars+raw_html",
        "--to", "docx",
        "--output", output_docx,
        "--wrap=none",
        "--columns=999",
        "-M", "east_asian_line_breaks=true",
        "--embed-resources"
    ]
    if reference_doc:
        cmd += ["--reference-doc", reference_doc]
    cmd.append(input_md)

    subprocess.run(cmd, check=True)

def clean_table_cells(docx_path):
    """
    遍历 DOCX 中的所有表格单元：
    - 若内容是 $...$ 或 $$...$$，去除外层 $，保留内部文本；
    - 若内容以 "=" 开头，标记为公式文本，使用等宽字体。
    注：不改变 Pandoc 生成的表格结构与样式。
    """
    doc = Document(docx_path)
    math_inline_re = re.compile(r'^\s\$(.+)\$\s*$')
    math_block_re = re.compile(r'^\s\$\$(.+)\$\$\s*$')

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # 合并同一单元的多个段落文本
                full_text = "\n".join(p.text for p in cell.paragraphs).strip()
                # 清理 $...$ 包裹
                cleaned = None
                m1 = math_inline_re.match(full_text)
                m2 = math_block_re.match(full_text)
                if m1:
                    cleaned = m1.group(1).strip()
                elif m2:
                    cleaned = m2.group(1).strip()

                if cleaned is not None:
                    # 覆盖为清洁文本
                    for p in cell.paragraphs:
                        # 清空现有 runs
                        for r in p.runs:
                            r.clear()
                        p.text = cleaned
                    continue  # 下一个单元

                # 若是 "=..."，设置等宽字体便于识别公式文本
                if full_text.startswith("="):
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.font.name = "Consolas"
                            r.font.size = Pt(10)
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.save(docx_path)

def main():
    parser = argparse.ArgumentParser(
        description="Convert Markdown to Word .docx with tables preserved as Word tables and formula text kept."
    )
    parser.add_argument("input_md", help="Input markdown file, e.g., 2.md")
    parser.add_argument("output_docx", nargs="?", help="Output docx file, e.g., 2.docx")
    parser.add_argument("--reference-doc", help="Optional reference DOCX to control styles", default=None)
    args = parser.parse_args()

    input_md = args.input_md
    if not os.path.isfile(input_md):
        print(f"Input file not found: {input_md}")
        sys.exit(1)

    output_docx = args.output_docx if args.output_docx else os.path.splitext(input_md)[0] + ".docx"

    if not check_pandoc():
        print("Error: pandoc not found. Please install pandoc from https://pandoc.org/installing.html")
        sys.exit(1)

    pre_md = preprocess_markdown(input_md)
    try:
        convert_md_to_docx(pre_md, output_docx, reference_doc=args.reference_doc)
        # 二次处理：清理表格单元中的 $...$ 与 =... 文本
        clean_table_cells(output_docx)
        print(f"Converted: {input_md} -> {output_docx}")
    except subprocess.CalledProcessError as e:
        print("Pandoc conversion failed.")
        print(e)
        sys.exit(1)
    finally:
        try:
            os.unlink(pre_md)
        except Exception:
            pass

if __name__ == "__main__":
    main()
